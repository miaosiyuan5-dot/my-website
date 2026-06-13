from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

OUTPUT_DIR = "/home/user/my-website/ted_scripts"

SCRIPTS = [
    {
        "filename": "01_The_Power_of_Small_Habits.docx",
        "title": "The Power of Small Habits",
        "subtitle": "Topic: Personal Growth | Level: B2 | ~6 min",
        "tip": "Shadowing tip: Focus on the stressed words in each sentence. Record yourself and compare.",
        "sections": [
            ("OPENING", """Imagine waking up every morning feeling like you are already behind. You hit the snooze button. You skip breakfast. You rush out the door. And by ten o'clock, you already feel exhausted.

Sound familiar?

That was me, three years ago. I was busy, stressed, and going nowhere fast. And then one small thing changed everything."""),
            ("THE STORY", """It started with a glass of water.

A friend told me, "Every morning, before you do anything else, drink a full glass of water." That's it. No big plan. No expensive program. Just water.

I thought it was too simple to matter. But I tried it. And something interesting happened.

Drinking that glass of water each morning gave me a tiny feeling of success. I had done something right before the day had even started. That small win made me want another one.

So I added five minutes of stretching. Then ten minutes of reading. Then I stopped checking my phone before breakfast.

Within two months, my mornings looked completely different — not because I had made one big decision, but because I had stacked many small ones."""),
            ("THE SCIENCE", """This is what scientists call a habit loop. Every habit has three parts: a cue, a routine, and a reward.

The cue is what triggers the behavior. The routine is the behavior itself. And the reward is what your brain gets at the end.

Here is the key insight: you cannot easily remove a bad habit, but you can replace it.

Research from Duke University found that about 40 percent of our daily actions are habits — not decisions. That means nearly half of what you do today, you are doing on autopilot."""),
            ("THE CHALLENGE", """Now, I know what some of you are thinking: "I've tried to build habits before. I always fail after two weeks."

The problem is not willpower. The problem is the plan.

Research shows that the most effective way to start a new habit is to make it so easy that you cannot say no. Want to start exercising? Do not begin with a one-hour gym session. Begin with two minutes of movement.

Start so small it feels almost embarrassing. Because a habit you do consistently — even a tiny one — beats a perfect habit that you quit."""),
            ("CLOSING", """I want to leave you with this thought:

You do not rise to the level of your goals. You fall to the level of your systems.

So do not ask yourself, "What do I want to achieve?" Ask yourself, "Who do I want to become?" And then ask, "What does that kind of person do every single day?"

Start with one habit. Make it tiny. Do it every day. And watch what happens over time.

The life you want is built one small choice at a time.

Thank you."""),
        ],
    },
    {
        "filename": "02_Why_Sleep_Is_Your_Superpower.docx",
        "title": "Why Sleep Is Your Superpower",
        "subtitle": "Topic: Health & Science | Level: B2 | ~6 min",
        "tip": "Shadowing tip: Pay attention to the rise and fall of intonation in questions.",
        "sections": [
            ("OPENING", """I want to ask you a simple question. When was the last time you woke up without an alarm, feeling completely refreshed?

If you have to think hard about that answer, you are not alone.

We live in a world that celebrates being busy. People wear their lack of sleep like a badge of honor. "I only slept four hours." "I'll sleep when I'm dead."

But here is the truth: if you keep sleeping too little, death comes sooner than you think."""),
            ("WHAT SLEEP DOES", """While you sleep, your brain is not resting. It is working.

During deep sleep, your brain flushes out toxic waste products — including proteins linked to Alzheimer's disease. Think of it as a dishwasher for your brain. Every night, it cleans out the mess from the day.

Sleep also strengthens memory. When you learn something new, your brain stores it temporarily. But during sleep, it transfers that information into long-term memory. This is why pulling an all-nighter before an exam is one of the worst strategies possible."""),
            ("THE COST OF BAD SLEEP", """The consequences of poor sleep are serious.

After just one night of sleeping only four to five hours, your immune system drops by 70 percent. Your risk of heart attack increases significantly. Your ability to regulate emotions collapses. You become more anxious, more aggressive, and less creative.

And the scary part? Most sleep-deprived people do not know how impaired they are. We lose the ability to accurately judge our own performance. We think we are fine. We are not."""),
            ("WHAT YOU CAN DO", """So what can you do?

First, protect your sleep schedule. Go to bed and wake up at the same time every day — even on weekends. Your brain runs on a biological clock, and consistency is the key.

Second, keep your bedroom cool and dark. Your body needs to drop its temperature to fall asleep. A room that is slightly cool — around 18 degrees Celsius — is ideal.

Third, avoid screens for at least 30 minutes before bed. The blue light from phones tricks your brain into thinking it is still daytime."""),
            ("CLOSING", """I want to leave you with a challenge.

Tonight, go to bed 30 minutes earlier than usual. Not tomorrow. Tonight.

Think of sleep not as a luxury, but as the foundation of everything else you want to do. Your productivity, your creativity, your happiness — they all depend on it.

The most powerful thing you can do for your health, your relationships, and your work might just be to close your eyes a little earlier.

Good night — and I mean that literally.

Thank you."""),
        ],
    },
    {
        "filename": "03_The_Gift_of_Failure.docx",
        "title": "The Gift of Failure",
        "subtitle": "Topic: Mindset & Education | Level: B2 | ~6 min",
        "tip": "Shadowing tip: Notice how the speaker uses pauses for dramatic effect. Try to copy the rhythm.",
        "sections": [
            ("OPENING", """When I was twelve years old, I failed my first piano recital.

I had practiced for months. I knew every note. But when I walked on stage, looked at the audience, and sat down at that piano — my mind went completely blank.

I stumbled through the piece. I hit wrong notes. And at the end, there was polite, quiet applause.

I was devastated. I went home and told my mother I would never play piano again.

She smiled and said, "Good. Now the real learning begins."

I had no idea what she meant. But twenty years later, I finally do."""),
            ("HOW WE SEE FAILURE", """We have a complicated relationship with failure.

On one hand, we say things like "learn from your mistakes" and "failure is the best teacher." But on the other hand, we design our schools, our workplaces, and our lives to avoid failure at all costs.

Students are punished for wrong answers. Employees fear making mistakes. We filter our lives on social media to show only the highlights.

The result? A generation that is terrified of trying anything new."""),
            ("WHAT RESEARCH SHOWS", """Here is what the research actually shows.

Stanford psychologist Carol Dweck spent decades studying how people respond to challenges. She found two types of thinking. People with a fixed mindset believe their abilities are set — they are either smart or they are not. So when they fail, they see it as proof that they are not good enough.

But people with a growth mindset believe abilities can be developed through effort. When they fail, they see it as information — a signal that they need to try a different approach.

The good news? Mindset can be changed."""),
            ("THE REAL LESSON", """The most successful people in the world are not the ones who never failed. They are the ones who failed the most — and kept going.

Thomas Edison tried over a thousand materials before finding the right one for the light bulb. When a reporter asked him how it felt to fail a thousand times, Edison said, "I didn't fail a thousand times. I found a thousand ways that don't work."

That is not just a clever answer. That is a completely different way of seeing the world."""),
            ("CLOSING", """So here is what I want you to do.

Find something you have been avoiding because you are afraid to fail. A language. An instrument. A conversation you keep putting off.

Start. Do it badly. Let yourself be a beginner.

Because here is the truth: every expert was once a beginner. Every master was once a disaster. The only difference is they did not quit after the first wrong note.

Failure is not the opposite of success. It is part of the path.

Thank you."""),
        ],
    },
    {
        "filename": "04_How_Kindness_Changes_Your_Brain.docx",
        "title": "How Kindness Changes Your Brain",
        "subtitle": "Topic: Psychology & Wellbeing | Level: B2 | ~6 min",
        "tip": "Shadowing tip: This talk has warm, gentle tone. Practice speaking with a soft, friendly voice.",
        "sections": [
            ("OPENING", """A few years ago, I was having a terrible day.

My train was late. I spilled coffee on my shirt. I missed an important meeting. By noon, I felt like the universe was personally against me.

Then a stranger held the door open for me. She smiled and said, "Hope your day gets better."

That was it. Ten words from someone I had never met.

But for the rest of that afternoon, I felt lighter. And I started wondering: why does kindness feel so powerful — even when it comes from a stranger?"""),
            ("THE SCIENCE OF KINDNESS", """It turns out, kindness is not just nice. It is neurological.

When you do something kind for someone else, your brain releases oxytocin — sometimes called the "love hormone." It also releases serotonin and dopamine, the same chemicals involved in happiness and reward.

But here is the fascinating part: the kindness effect works in three directions. The person who gives kindness feels it. The person who receives it feels it. And even people who simply witness an act of kindness experience a boost in their own well-being.

Scientists call this "elevation." Witnessing goodness literally lifts us up."""),
            ("THE RIPPLE EFFECT", """Kindness also spreads.

A study from the University of California found that when people receive a kind act, they are significantly more likely to be kind to someone else — often someone completely different. Researchers call this "paying it forward."

One act of kindness can trigger a chain reaction. A stranger's smile leads to a held door. A held door leads to a "thank you." A "thank you" leads to a better conversation. And slowly, the mood of an entire community shifts."""),
            ("WHAT STOPS US", """So if kindness feels good and spreads easily, why don't we do it more?

Often, we are simply too distracted. We are looking at our phones. We are thinking about our to-do lists. We miss the opportunities for connection that are right in front of us.

Sometimes we feel too rushed. We tell ourselves, "I'll help next time." But next time never comes.

And sometimes, honestly, we feel that kindness makes us look weak. In a competitive world, vulnerability can feel dangerous.

But research consistently shows the opposite: kind people are seen as stronger, more trustworthy, and more capable — not weaker."""),
            ("CLOSING", """I want to give you a very small challenge.

In the next 24 hours, do one act of kindness that costs you nothing but attention. Hold the door. Give a genuine compliment. Put your phone down and listen — really listen — to someone who needs to talk.

Notice how you feel afterward.

Because kindness is not just good for the world. It is good for you. It rewires your brain. It strengthens your immune system. It lowers your blood pressure.

In a world that often feels divided and difficult, kindness is not weakness.

It might just be the bravest thing you can do.

Thank you."""),
        ],
    },
    {
        "filename": "05_The_Danger_of_Comparing_Yourself_to_Others.docx",
        "title": "The Danger of Comparing Yourself to Others",
        "subtitle": "Topic: Mental Health & Social Media | Level: B2 | ~6 min",
        "tip": "Shadowing tip: This talk has moments of humor. Practice delivering the funny lines with a light, relaxed tone.",
        "sections": [
            ("OPENING", """I have a confession to make.

Every morning, before I even get out of bed, I check Instagram. I look at someone's perfect breakfast. Someone else's vacation photos. A friend's promotion announcement. And within five minutes, before my day has even started, I already feel behind.

Does anyone else do this? Just me?

I didn't think so.

We live in the age of comparison. And it is quietly making us miserable."""),
            ("WHY WE COMPARE", """Comparing ourselves to others is actually a very human instinct.

Psychologist Leon Festinger called it "social comparison theory." He argued that humans naturally evaluate themselves by looking at others. In prehistoric times, this was useful. It helped us understand where we stood in the group, who to learn from, who to compete with.

But social media has taken this ancient instinct and turned it up to an impossible volume.

We are no longer comparing ourselves to our neighbors. We are comparing ourselves to the most successful, most attractive, most curated version of millions of people — all at once, all day long."""),
            ("THE HIGHLIGHT REEL PROBLEM", """Here is the thing about social media: everyone posts their best moments.

No one posts a photo of themselves crying in a parking lot. No one shares the argument they had with their partner this morning. No one announces that they stayed in bed all day because they felt hopeless.

What we see online is a highlight reel. But we compare it to our behind-the-scenes footage — all our doubts, our bad days, our private struggles.

It is a deeply unfair comparison. And we make it dozens of times a day."""),
            ("A BETTER WAY", """So what is the alternative?

Researchers suggest shifting from social comparison to what they call "temporal comparison" — comparing yourself not to others, but to your past self.

Ask: Am I better than I was a year ago? Am I kinder? Am I more skilled? Am I learning?

This kind of comparison is motivating rather than demoralizing. Because the only competitor who actually matters is the person you were yesterday."""),
            ("CLOSING", """I am not saying delete your social media. I am not saying ignore the world around you.

I am saying: be careful about the stories you tell yourself when you scroll.

The person with the perfect life online may be struggling in ways you cannot see. And the life you are living — with all its imperfections, all its ordinary moments — may be more valuable than you realize.

Run your own race. At your own pace. On your own terms.

Because the only finish line that matters is the one you set for yourself.

Thank you."""),
        ],
    },
    {
        "filename": "06_Why_Boredom_Is_Good_for_You.docx",
        "title": "Why Boredom Is Good for You",
        "subtitle": "Topic: Creativity & Technology | Level: B2 | ~6 min",
        "tip": "Shadowing tip: Practice the rhetorical questions — let your voice go up slightly at the end.",
        "sections": [
            ("OPENING", """When was the last time you were truly bored?

Not waiting-for-the-bus bored. Not bored-while-checking-your-phone bored. I mean genuinely, deeply, with-nothing-to-do bored.

For most of us, that moment barely exists anymore. Every idle second is filled — with a scroll, a podcast, a notification. We have built an entire civilization designed to eliminate boredom.

And I think that is a serious problem."""),
            ("WHAT BOREDOM ACTUALLY DOES", """Boredom is uncomfortable. We all know that feeling. Your mind wanders. You feel restless. You want something — anything — to fill the silence.

But that wandering? That restlessness? That is your brain entering what neuroscientists call the "default mode network."

When your mind is free to roam without a specific task, it does something remarkable. It starts to make connections. It replays memories, imagines the future, and combines ideas in new ways.

In other words: boredom is when creativity happens."""),
            ("THE EVIDENCE", """Consider some of history's most creative moments.

Newton was sitting in a garden, doing nothing, when he observed a falling apple and began thinking about gravity.

Archimedes was relaxing in a bath when he suddenly understood how to measure volume.

J.K. Rowling was sitting on a delayed train with nothing to do when the idea for Harry Potter arrived in her mind.

These are not coincidences. These are minds that were given space — and used it."""),
            ("WHAT WE LOSE WHEN WE ARE NEVER BORED", """When we fill every moment with stimulation, we lose something essential.

We lose the ability to sit with our own thoughts. We lose patience. We lose the capacity for deep focus — because we have trained our brains to constantly expect something new.

Children who are never allowed to be bored often struggle to entertain themselves. Adults who are never offline often find they cannot concentrate for more than a few minutes.

We have outsourced our inner life to our devices. And our inner life is getting weaker."""),
            ("CLOSING", """Here is my challenge to you.

Once a day, for just ten minutes, do nothing. No phone. No music. No podcast. Just sit. Walk without headphones. Eat without a screen.

It will feel uncomfortable at first. Your brain will resist. But stay with it.

Because in that discomfort, in that quiet space between tasks — that is where your best ideas live.

Boredom is not the enemy of productivity. It is the source of it.

Let yourself be bored.

Thank you."""),
        ],
    },
    {
        "filename": "07_The_Language_of_Body_Language.docx",
        "title": "The Language of Body Language",
        "subtitle": "Topic: Communication & Psychology | Level: B2 | ~6 min",
        "tip": "Shadowing tip: This topic is about non-verbal cues — as you read aloud, try to match your posture to the content.",
        "sections": [
            ("OPENING", """Before you say a single word, people have already formed an impression of you.

They have noticed how you walked into the room. Whether you made eye contact. How you are holding your arms. Whether you seem open or closed, confident or nervous.

Research suggests that in a first impression, your words account for only about seven percent of the message. Your tone of voice accounts for 38 percent. And your body language — your posture, your gestures, your expression — accounts for 55 percent.

Your body is speaking. The question is: what is it saying?"""),
            ("WHAT YOUR BODY REVEALS", """Body language is a two-way conversation between your mind and the people around you.

When you cross your arms, you might feel defensive or cold — and others read it that way. When you lean forward, you signal interest and engagement. When you maintain steady eye contact, you communicate confidence and honesty.

But here is what most people don't realize: your body language affects not just how others see you — it affects how you see yourself.

Social psychologist Amy Cuddy found that holding a "power pose" for just two minutes — standing tall, shoulders back, taking up space — changes the level of hormones in your body. It raises testosterone and lowers cortisol. You feel more confident. And that confidence shows."""),
            ("READING OTHERS", """Understanding body language also helps you understand people around you.

Is someone truly listening, or just waiting to speak? Watch their eyes. If they are genuinely engaged, their pupils slightly dilate and they lean toward you. If they are bored or distracted, they will look around, shift their weight, or point their feet toward the nearest exit — literally.

Eye contact is particularly powerful. Sustained, comfortable eye contact signals trust and connection. Avoiding eye contact can signal discomfort, dishonesty, or simply shyness. Context matters enormously."""),
            ("COMMON MISTAKES", """There are common body language mistakes that undermine even the most well-prepared speakers.

Touching your face often signals anxiety or deception — even when you feel neither. Slouching communicates low energy or low confidence. Speaking with your hands hidden — in your pockets or behind your back — creates a sense of distance and reduces trust.

And the biggest mistake? Not making enough eye contact. When you look down at your notes constantly, you lose the human connection that makes communication powerful."""),
            ("CLOSING", """Here is something I want you to try in your next important conversation or presentation.

Before you begin, take a breath. Stand up straight — not stiffly, but open. Relax your shoulders. Make eye contact with one person for a full sentence before moving to the next.

Notice how the room responds differently. Notice how you feel differently.

Your body is your first language. Learn to speak it well.

Because sometimes, the most powerful thing you can say has nothing to do with words.

Thank you."""),
        ],
    },
    {
        "filename": "08_The_Art_of_Saying_No.docx",
        "title": "The Art of Saying No",
        "subtitle": "Topic: Wellbeing & Productivity | Level: B2 | ~6 min",
        "tip": "Shadowing tip: Notice the direct, confident tone. Practice speaking each sentence firmly and clearly.",
        "sections": [
            ("OPENING", """I used to say yes to everything.

Every invitation. Every request. Every "can you just quickly help me with this?" I said yes because I wanted to be helpful. I said yes because I was afraid of disappointing people. I said yes because somewhere, I believed that being busy meant being important.

And then I burned out.

I was exhausted, overwhelmed, and resentful — not of the people who asked, but of myself for agreeing.

That was the day I learned that "no" is a complete sentence."""),
            ("WHY WE STRUGGLE TO SAY NO", """Saying no is hard. There are real reasons for this.

We are social creatures. We want to be liked. We fear conflict. We worry that saying no will damage relationships or opportunities.

In many cultures, saying no is considered rude or selfish. We are raised to be agreeable, cooperative, accommodating. And those are not bad values. But taken too far, they leave us with no time for the things that actually matter to us."""),
            ("THE COST OF ALWAYS SAYING YES", """Every time you say yes to something you do not want — or cannot afford — to do, you are saying no to something else.

You say yes to an extra project at work. You say no to dinner with your family. You say yes to a social obligation you dread. You say no to the rest you desperately need. You say yes to other people's priorities. You say no to your own.

Warren Buffett once said that the difference between successful people and very successful people is that very successful people say no to almost everything.

No is not a rejection. No is a direction."""),
            ("HOW TO SAY NO WELL", """So how do you say no without damaging relationships?

First, you do not need to apologize or over-explain. "I won't be able to make that work" is a complete and honest answer. You don't owe anyone a detailed justification for your time.

Second, you can be warm without being weak. "I appreciate you thinking of me, but no" is kind and clear at the same time.

Third, delay if you need to. Instead of an automatic yes, try "Let me check my schedule and get back to you." This gives you space to make a real decision rather than a reflexive one."""),
            ("CLOSING", """Every no to the wrong thing is a yes to the right thing.

Every time you protect your time and energy, you have more to give to the people and projects that truly deserve it.

Saying no is not selfish. It is honest. It is respectful — to yourself, and to the people you say it to. Because a reluctant yes is worse than a clear no.

Your time is the one thing you can never get back. Spend it like it matters.

Because it does.

Thank you."""),
        ],
    },
    {
        "filename": "09_What_Travel_Teaches_You_About_Yourself.docx",
        "title": "What Travel Teaches You About Yourself",
        "subtitle": "Topic: Culture & Self-Discovery | Level: B2 | ~6 min",
        "tip": "Shadowing tip: This is a narrative, storytelling talk. Let your voice be expressive and varied in pace.",
        "sections": [
            ("OPENING", """The first time I traveled alone, I was completely lost — in every sense of the word.

I was in a city where I did not speak the language. My phone had no signal. My map was wrong. And I had no idea where my hotel was.

And here is the strange part: it was one of the best experiences of my life.

Not because getting lost is fun. But because of what I discovered when I had no choice but to figure it out."""),
            ("TRAVEL AS A MIRROR", """Travel has a way of showing you who you really are — not the version of yourself you present to people who know you, but the real version, under pressure.

Are you patient when things go wrong? Are you flexible when your plan falls apart? Are you willing to ask for help from a stranger?

At home, we operate in our comfort zone. Everything is familiar. Our habits, our routines, our social roles — they are all in place. But travel strips those things away. And what's left is just you.

That can be uncomfortable. And that discomfort is exactly the point."""),
            ("WHAT YOU LEARN FROM OTHER CULTURES", """Every culture has solved the basic problems of human life differently.

How do you greet a stranger? How do you show respect? What do you eat, and with whom, and when? What do you celebrate? What do you mourn?

When you travel to a place where everything is done differently, you realize something profound: the way you grew up doing things is not the only way. It is one way, among many.

This is not a small realization. It is humbling and liberating at the same time. It makes you more open, more curious, and more compassionate."""),
            ("THE LESSONS THAT LAST", """The lessons from travel are not always comfortable.

You learn that communication is possible even without shared language. A smile, a gesture, a willingness to try — these cross any border.

You learn that most people, in most places, are kind. They want to help you find your way. They are proud of where they come from and happy to share it.

And you learn that home — whatever and wherever that is for you — is not a fixed place. It is a feeling. And that feeling can exist in many places, with many people, in many forms."""),
            ("CLOSING", """You don't have to travel to a distant country to experience this. A new neighborhood, a different community, a conversation with someone whose life looks nothing like yours — these are all forms of travel.

The point is to step outside the familiar. To be the person who does not know. To be the beginner, the visitor, the one who asks questions.

Because every time you leave your comfort zone, you come back a little larger than you were before.

Go somewhere new. Talk to a stranger. Get a little lost.

You might just find yourself.

Thank you."""),
        ],
    },
    {
        "filename": "10_The_Power_of_Listening.docx",
        "title": "The Power of Listening",
        "subtitle": "Topic: Communication & Relationships | Level: B2 | ~6 min",
        "tip": "Shadowing tip: This talk is calm and reflective. Practice speaking slowly, with long pauses at the dots.",
        "sections": [
            ("OPENING", """We live in a world full of noise.

Everyone is talking. Social media is shouting. News channels are competing for your attention 24 hours a day. And in the middle of all this noise, something very quiet and very important is disappearing.

The ability to truly listen.

Not just to hear the sounds. But to be fully present with another person. To understand not just their words, but what those words mean to them.

I believe listening may be the most underrated skill of the twenty-first century."""),
            ("THE DIFFERENCE BETWEEN HEARING AND LISTENING", """Hearing is passive. Sound enters your ears and your brain processes it automatically. You do not choose to hear.

But listening is active. Listening requires attention, patience, and intention. When you listen, you are not just processing words. You are reading tone, noticing what is not said, and making the other person feel that they matter.

Research shows that most people speak at about 130 words per minute, but we can process information at about 400 words per minute. That gap — that extra mental space — is usually where we stop listening and start planning what to say next."""),
            ("WHY WE ARE POOR LISTENERS", """We are poor listeners for many reasons.

We are distracted. Phones, notifications, background noise — they all compete for our attention. Even in a face-to-face conversation, part of our mind is somewhere else.

We are impatient. We think we know what the other person is going to say, so we finish their sentences — in our heads or out loud — and move on.

And we are self-focused. We listen to respond, not to understand. We wait for our turn, and while waiting, we miss everything the other person is actually trying to tell us."""),
            ("HOW LISTENING CHANGES RELATIONSHIPS", """The impact of listening — real listening — is extraordinary.

Studies consistently show that people feel more understood, more valued, and more connected when someone truly listens to them. They also feel that the listener is more intelligent and more trustworthy.

And here is the paradox: in a conversation, the person who talks the least often has the most influence. People trust those who listen. They open up to them. They want to be around them.

The best leaders, the best therapists, the best friends — they are almost always excellent listeners."""),
            ("CLOSING", """Here is a simple practice.

In your next conversation, try this: do not plan your response while the other person is speaking. Simply listen. When they finish, take a breath. Then respond.

You will be surprised how much you missed before. You will be surprised how much more connected the conversation feels.

In a world that rewards those who speak the loudest, choose to be someone who listens the deepest.

Because the greatest gift you can give another person is not advice, not money, not time.

It is your full, undivided attention.

Thank you."""),
        ],
    },
]


def add_title_section(doc, title, subtitle, tip):
    # Title
    t = doc.add_heading(title, level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in t.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    # Subtitle
    sub = doc.add_paragraph(subtitle)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in sub.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x77)

    doc.add_paragraph()

    # Tip box
    tip_para = doc.add_paragraph()
    tip_run = tip_para.add_run("💡 " + tip)
    tip_run.font.size = Pt(10)
    tip_run.font.italic = True
    tip_run.font.color.rgb = RGBColor(0x2E, 0x86, 0xAB)

    doc.add_paragraph()


def add_section(doc, heading, content):
    h = doc.add_heading(heading, level=2)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x2E, 0x86, 0xAB)

    paragraphs = content.strip().split("\n\n")
    for para_text in paragraphs:
        p = doc.add_paragraph(para_text.strip())
        p.paragraph_format.space_after = Pt(8)
        for run in p.runs:
            run.font.size = Pt(12)

    doc.add_paragraph()


def create_doc(script):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(12)

    add_title_section(doc, script["title"], script["subtitle"], script["tip"])

    for heading, content in script["sections"]:
        add_section(doc, heading, content)

    # Footer note
    doc.add_paragraph()
    footer_para = doc.add_paragraph("— End of Script —")
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer_para.runs:
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    path = os.path.join(OUTPUT_DIR, script["filename"])
    doc.save(path)
    print(f"✓ Saved: {script['filename']}")


for script in SCRIPTS:
    create_doc(script)

print("\nAll 10 scripts generated successfully!")
